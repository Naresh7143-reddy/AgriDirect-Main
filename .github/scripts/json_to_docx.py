#!/usr/bin/env python3
"""
Convert JSON test results to Word document (.docx)

Usage: python3 json_to_docx.py <input.json> "<title>" <output.docx>
"""

import json
import sys
import os
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx is not installed. Please install it first:")
    print("  pip install python-docx")
    sys.exit(1)


def set_cell_background(cell, fill):
    """Set cell background color"""
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)


def add_status_color(cell, status):
    """Add color to cell based on status"""
    colors = {
        "passed": "31A24C",
        "failed": "D32F2F",
        "skipped": "FFA500",
    }
    if status and status.lower() in colors:
        set_cell_background(cell, colors[status.lower()])


def get_status_emoji(status):
    """Get emoji for status"""
    emojis = {"passed": "✓", "failed": "✗", "skipped": "⊘"}
    return emojis.get(status.lower(), "•") if status else "•"


def main():
    if len(sys.argv) < 3:
        print("Usage: python3 json_to_docx.py <input.json> <title> <output.docx>")
        sys.exit(1)

    input_file = sys.argv[1]
    title = sys.argv[2] if len(sys.argv) > 2 else "Test Report"
    output_file = sys.argv[3] if len(sys.argv) > 3 else "report.docx"

    # Parse JSON
    test_data = {
        "suite": title,
        "total": 0,
        "passed": 0,
        "failed": 0,
        "skipped": 0,
        "cases": [],
    }

    if os.path.exists(input_file):
        try:
            with open(input_file, "r") as f:
                loaded_data = json.load(f)
                test_data.update(loaded_data)
        except Exception as e:
            print(f"Warning: Could not parse {input_file}: {e}")

    total = test_data.get("total", 0)
    passed = test_data.get("passed", 0)
    failed = test_data.get("failed", 0)
    skipped = test_data.get("skipped", 0)
    cases = test_data.get("cases", [])
    success_rate = (passed / total * 100) if total > 0 else 0

    # Create document
    doc = Document()

    # Add title
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(test_data["suite"])
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(54, 96, 146)  # Blue

    # Add summary section
    doc.add_paragraph("Summary", style="Heading 2")

    # Summary table
    summary_table = doc.add_table(rows=2, cols=5)
    summary_table.style = "Light Grid Accent 1"

    # Header row
    header_cells = summary_table.rows[0].cells
    header_cells[0].text = "Total Tests"
    header_cells[1].text = "Passed"
    header_cells[2].text = "Failed"
    header_cells[3].text = "Skipped"
    header_cells[4].text = "Success Rate"

    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, "366092")

    # Data row
    data_cells = summary_table.rows[1].cells
    data_cells[0].text = str(total)
    data_cells[1].text = str(passed)
    data_cells[2].text = str(failed)
    data_cells[3].text = str(skipped)
    data_cells[4].text = f"{success_rate:.1f}%"

    for i, cell in enumerate(data_cells):
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(12)

        if i == 1:  # Passed column
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(49, 162, 76)
        elif i == 2:  # Failed column
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(211, 47, 47)
        elif i == 3:  # Skipped column
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 165, 0)

    # Add test cases section if there are any
    if cases:
        doc.add_paragraph()
        doc.add_paragraph("Test Cases", style="Heading 2")

        # Test cases table
        test_table = doc.add_table(rows=len(cases) + 1, cols=5)
        test_table.style = "Light Grid Accent 1"

        # Header row
        header_row = test_table.rows[0]
        headers = ["#", "Test Case", "File", "Status", "Duration (ms)"]
        for i, header_text in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header_text
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            set_cell_background(cell, "366092")
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Data rows
        for row_idx, case in enumerate(cases):
            row = test_table.rows[row_idx + 1]
            row.cells[0].text = str(row_idx + 1)
            row.cells[1].text = case.get("name", "Unknown")
            row.cells[2].text = case.get("file", "N/A")
            status = case.get("status", "unknown")
            row.cells[3].text = f"{get_status_emoji(status)} {status.upper()}"
            row.cells[4].text = str(case.get("duration", 0))

            # Set alignment
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Set status color
            add_status_color(row.cells[3], status)

    # Add footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph(
        f"Generated on {datetime.now().isoformat()}"
    )
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.runs[0].font.size = Pt(9)
    footer_para.runs[0].font.italic = True
    footer_para.runs[0].font.color.rgb = RGBColor(153, 153, 153)

    # Create output directory
    output_dir = os.path.dirname(output_file)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    # Save document
    doc.save(output_file)
    print(f"✓ Word document generated: {output_file}")
    print(f"  - Total tests: {total}")
    print(f"  - Passed: {passed} ({success_rate:.1f}%)")
    print(f"  - Failed: {failed}")
    print(f"  - Skipped: {skipped}")


if __name__ == "__main__":
    main()
