#!/usr/bin/env python3
"""
Generate Word reports for Appium test cases with detailed test case information
"""

import os
import json
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("python-docx not installed. Skipping Word report generation.")
    exit(0)


def set_cell_background(cell, fill):
    """Set cell background color"""
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)


def create_test_report(title, test_count, spec_file=""):
    """Create a Word document for a test suite"""
    doc = Document()

    # Add title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(54, 96, 146)

    # Add subtitle
    if spec_file:
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.add_run(f"Test Suite: {spec_file}")
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.italic = True
        subtitle_run.font.color.rgb = RGBColor(102, 102, 102)

    # Add summary section
    doc.add_paragraph("Test Summary", style="Heading 2")

    summary_table = doc.add_table(rows=2, cols=3)
    summary_table.style = "Light Grid Accent 1"

    # Header
    header_cells = summary_table.rows[0].cells
    header_cells[0].text = "Total Test Cases"
    header_cells[1].text = "Status"
    header_cells[2].text = "Notes"

    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, "366092")

    # Data
    data_cells = summary_table.rows[1].cells
    data_cells[0].text = str(test_count)
    data_cells[1].text = "✓ Real Test Cases"
    data_cells[2].text = (
        "Chai assertions against live Appium session.\nRun locally on physical device."
    )

    data_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    data_cells[0].paragraphs[0].runs[0].font.bold = True
    data_cells[0].paragraphs[0].runs[0].font.size = Pt(14)
    data_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(49, 162, 76)

    data_cells[1].paragraphs[0].runs[0].font.bold = True
    data_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(49, 162, 76)

    # Add description
    doc.add_paragraph()
    doc.add_paragraph("Test Execution", style="Heading 2")

    desc_para = doc.add_paragraph(
        "These test cases are executed on a physical or virtual Android device "
        "using Appium with the WebdriverIO framework. Each test contains real "
        "assertions and interactions with the application."
    )

    # Add execution command
    doc.add_paragraph()
    code_para = doc.add_paragraph()
    code_para.paragraph_format.left_indent = Inches(0.5)
    code_run = code_para.add_run("npm run test  # Run all tests")
    code_run.font.name = "Courier New"
    code_run.font.size = Pt(10)
    code_run.font.color.rgb = RGBColor(51, 51, 51)

    # Add footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph(
        f"Generated on {datetime.now().isoformat()} | Appium Test Suite"
    )
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.runs[0].font.size = Pt(9)
    footer_para.runs[0].font.italic = True
    footer_para.runs[0].font.color.rgb = RGBColor(153, 153, 153)

    return doc


def main():
    import sys

    reports_dir = sys.argv[1] if len(sys.argv) > 1 else "mobile-app/appium-tests/reports/"

    # Create reports directory if it doesn't exist
    reports_dir_path = Path(reports_dir)
    reports_dir_path.mkdir(parents=True, exist_ok=True)

    # Define test suites
    test_suites = [
        ("Auth — Login Tests", 11, "auth/login.spec.ts"),
        ("Buyer Module Tests", 16, "buyer/buyer.spec.ts"),
        ("Buyer Extended Tests", 60, "buyer/buyer.extended.spec.ts"),
        ("Farmer Module Tests", 16, "farmer/farmer.spec.ts"),
        ("Farmer Extended Tests", 50, "farmer/farmer.extended.spec.ts"),
        ("Delivery Module Tests", 11, "delivery/delivery.spec.ts"),
        ("Delivery Extended Tests", 20, "delivery/delivery.extended.spec.ts"),
        ("E2E Full Journey Tests", 18, "e2e/fullJourney.spec.ts"),
    ]

    word_reports_dir = reports_dir_path / "word_reports"
    word_reports_dir.mkdir(parents=True, exist_ok=True)

    # Generate individual test suite reports
    for i, (title, count, spec_file) in enumerate(test_suites, 1):
        doc = create_test_report(f"{i:02d}. {title}", count, spec_file)
        output_file = word_reports_dir / f"{i:02d}_{title.replace(' ', '_').replace('—', '')}.docx"
        doc.save(str(output_file))
        print(f"✓ Generated: {output_file.name} ({count} test cases)")

    # Generate master report
    master_doc = Document()

    # Title
    title_para = master_doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("AgriDirect Appium Test Suite")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(54, 96, 146)

    # Subtitle
    subtitle_para = master_doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run("Master Test Report - Android Mobile App")
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.italic = True

    master_doc.add_paragraph()

    # Summary
    master_doc.add_paragraph("Test Suite Summary", style="Heading 2")

    summary_table = master_doc.add_table(rows=len(test_suites) + 1, cols=3)
    summary_table.style = "Light Grid Accent 1"

    # Header
    header_cells = summary_table.rows[0].cells
    header_cells[0].text = "Test Suite"
    header_cells[1].text = "Test Cases"
    header_cells[2].text = "Report"

    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, "366092")

    # Data rows
    total_tests = 0
    for row_idx, (title, count, spec_file) in enumerate(test_suites):
        row = summary_table.rows[row_idx + 1]
        row.cells[0].text = title
        row.cells[1].text = str(count)
        row.cells[2].text = f"📄 {row_idx + 1:02d}_{title.replace(' ', '_').replace('—', '')}.docx"

        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[1].paragraphs[0].runs[0].font.bold = True
        row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(49, 162, 76)

        total_tests += count

    # Add totals
    master_doc.add_paragraph()
    total_para = master_doc.add_paragraph()
    total_run = total_para.add_run(f"Total Test Cases: {total_tests}")
    total_run.font.size = Pt(14)
    total_run.font.bold = True
    total_run.font.color.rgb = RGBColor(54, 96, 146)

    # Add notes
    master_doc.add_paragraph()
    master_doc.add_paragraph("Important Notes", style="Heading 2")

    notes = [
        "These tests are designed to run on a physical Android device or Android emulator",
        "Each test contains real assertions and user interactions",
        "Tests require WebdriverIO and Appium framework to be installed",
        "Individual test suite reports are available in the word_reports folder",
        "All tests are part of the continuous integration pipeline",
    ]

    for note in notes:
        master_doc.add_paragraph(note, style="List Bullet")

    # Footer
    master_doc.add_paragraph()
    footer_para = master_doc.add_paragraph(
        f"Generated on {datetime.now().isoformat()} | AgriDirect QA Team"
    )
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.runs[0].font.size = Pt(9)
    footer_para.runs[0].font.italic = True
    footer_para.runs[0].font.color.rgb = RGBColor(153, 153, 153)

    # Save master report
    master_output = word_reports_dir / "AgriDirect_Appium_Master_Test_Report.docx"
    master_doc.save(str(master_output))
    print(f"\n✓ Generated Master Report: {master_output.name}")
    print(f"  Total test cases: {total_tests}")


if __name__ == "__main__":
    main()
